"""
文件服务
处理文件的上传、保存和读取
支持多种文档格式的解析和优化
"""

import os
import uuid
import aiofiles
from typing import Optional, Dict, Any, List
from fastapi import UploadFile
from datetime import datetime
import re


class FileService:
    """文件服务类"""

    def __init__(self, upload_dir: str = "./data/uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

        # 支持的文件类型
        self.supported_types = {
            ".txt": "text",
            ".md": "markdown",
            ".pdf": "pdf",
            ".docx": "word",
            ".doc": "word",
            ".html": "html",
            ".htm": "html",
            ".csv": "csv",
            ".json": "json"
        }

    async def save_upload_file(self, file: UploadFile) -> str:
        """
        保存上传的文件

        Args:
            file: 上传的文件

        Returns:
            文件路径
        """
        # 生成唯一文件名
        file_ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(self.upload_dir, unique_filename)

        # 保存文件
        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await file.read()
            await out_file.write(content)

        return file_path

    async def read_file(self, file_path: str) -> str:
        """
        读取文件内容

        Args:
            file_path: 文件路径

        Returns:
            文件内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == ".txt":
            return await self._read_text_file(file_path)
        elif file_ext == ".md":
            return await self._read_markdown_file(file_path)
        elif file_ext == ".pdf":
            return await self._read_pdf_file(file_path)
        elif file_ext in [".docx", ".doc"]:
            return await self._read_word_file(file_path)
        elif file_ext in [".html", ".htm"]:
            return await self._read_html_file(file_path)
        elif file_ext == ".csv":
            return await self._read_csv_file(file_path)
        elif file_ext == ".json":
            return await self._read_json_file(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")

    async def read_file_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        读取文件内容并返回元数据

        Args:
            file_path: 文件路径

        Returns:
            包含内容和元数据的字典
        """
        content = await self.read_file(file_path)
        file_info = self.get_file_info(file_path)

        # 提取文档结构
        structure = self._extract_structure(content, file_info["file_type"])

        return {
            "content": content,
            "metadata": file_info,
            "structure": structure,
            "word_count": len(content.split()),
            "char_count": len(content)
        }

    async def _read_text_file(self, file_path: str) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                    return await f.read()
            except UnicodeDecodeError:
                continue

        raise Exception(f"无法识别文件编码: {file_path}")

    async def _read_markdown_file(self, file_path: str) -> str:
        """读取Markdown文件"""
        content = await self._read_text_file(file_path)

        # 清理Markdown格式
        content = self._clean_markdown(content)

        return content

    async def _read_pdf_file(self, file_path: str) -> str:
        """读取PDF文件"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- 第 {i+1} 页 ---\n{page_text}")

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"读取PDF文件失败: {str(e)}")

    async def _read_word_file(self, file_path: str) -> str:
        """读取Word文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # 读取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 根据样式添加标题标记
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading', '')
                        text_parts.append(f"{'#' * int(level)} {paragraph.text}")
                    else:
                        text_parts.append(paragraph.text)

            # 读取表格
            for i, table in enumerate(doc.tables):
                text_parts.append(f"\n--- 表格 {i+1} ---")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"读取Word文件失败: {str(e)}")

    async def _read_html_file(self, file_path: str) -> str:
        """读取HTML文件"""
        try:
            from html.parser import HTMLParser

            class HTMLTextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text_parts = []
                    self.skip_tags = {'script', 'style', 'head'}
                    self.current_tag = None

                def handle_starttag(self, tag, attrs):
                    self.current_tag = tag
                    if tag in self.skip_tags:
                        self.skip = True

                def handle_endtag(self, tag):
                    if tag in self.skip_tags:
                        self.skip = False
                    self.current_tag = None

                def handle_data(self, data):
                    if hasattr(self, 'skip') and self.skip:
                        return
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)

                def get_text(self):
                    return ' '.join(self.text_parts)

            content = await self._read_text_file(file_path)

            extractor = HTMLTextExtractor()
            extractor.feed(content)

            return extractor.get_text()
        except Exception as e:
            raise Exception(f"读取HTML文件失败: {str(e)}")

    async def _read_csv_file(self, file_path: str) -> str:
        """读取CSV文件"""
        try:
            import csv

            content = await self._read_text_file(file_path)

            # 解析CSV
            reader = csv.reader(content.splitlines())
            rows = list(reader)

            # 转换为表格格式
            if len(rows) > 0:
                headers = rows[0]
                text_parts = [f"列名: {', '.join(headers)}"]
                text_parts.append("---")

                for i, row in enumerate(rows[1:], 1):
                    row_data = {headers[j]: row[j] for j in range(min(len(headers), len(row)))}
                    text_parts.append(f"行 {i}: {row_data}")

                return "\n".join(text_parts)
            else:
                return "空CSV文件"
        except Exception as e:
            raise Exception(f"读取CSV文件失败: {str(e)}")

    async def _read_json_file(self, file_path: str) -> str:
        """读取JSON文件"""
        try:
            import json

            content = await self._read_text_file(file_path)
            data = json.loads(content)

            # 格式化输出
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            raise Exception(f"读取JSON文件失败: {str(e)}")

    def _clean_markdown(self, content: str) -> str:
        """清理Markdown格式"""
        # 移除图片链接
        content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        # 移除链接但保留文本
        content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)
        # 移除HTML标签
        content = re.sub(r'<[^>]+>', '', content)
        # 清理多余的空行
        content = re.sub(r'\n{3,}', '\n\n', content)

        return content.strip()

    def _extract_structure(self, content: str, file_type: str) -> Dict[str, Any]:
        """提取文档结构"""
        structure = {
            "sections": [],
            "headings": [],
            "tables_count": 0,
            "lists_count": 0
        }

        # 提取标题
        if file_type == "markdown":
            headings = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
            structure["headings"] = [
                {"level": len(h[0]), "text": h[1]}
                for h in headings
            ]
        elif file_type == "word":
            # 从内容中提取标题
            lines = content.split('\n')
            for line in lines:
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    structure["headings"].append({"level": level, "text": text})

        # 统计表格
        structure["tables_count"] = content.count('--- 表格')

        # 统计列表
        structure["lists_count"] = len(re.findall(r'^[-*]\s+', content, re.MULTILINE))

        return structure

    def get_file_info(self, file_path: str) -> dict:
        """
        获取文件信息

        Args:
            file_path: 文件路径

        Returns:
            文件信息字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        stat = os.stat(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()

        return {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size": stat.st_size,
            "file_type": file_ext,
            "file_category": self.supported_types.get(file_ext, "unknown"),
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
        }

    def delete_file(self, file_path: str) -> bool:
        """
        删除文件

        Args:
            file_path: 文件路径

        Returns:
            是否删除成功
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False
        except Exception:
            return False

    def list_files(self, directory: Optional[str] = None) -> list:
        """
        列出目录中的文件

        Args:
            directory: 目录路径，默认为上传目录

        Returns:
            文件列表
        """
        target_dir = directory or self.upload_dir

        if not os.path.exists(target_dir):
            return []

        files = []
        for filename in os.listdir(target_dir):
            file_path = os.path.join(target_dir, filename)
            if os.path.isfile(file_path):
                files.append(self.get_file_info(file_path))

        return files

    def get_upload_path(self, filename: str) -> str:
        """获取上传文件的完整路径"""
        return os.path.join(self.upload_dir, filename)

    def file_exists(self, file_path: str) -> bool:
        """检查文件是否存在"""
        return os.path.exists(file_path)

    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型"""
        return list(self.supported_types.keys())
