"""
报告服务
处理报告的导出和格式转换
"""

import os
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class ReportService:
    """报告服务类"""

    def __init__(self, export_dir: str = "./data/exports"):
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)

    async def export_report(
        self,
        content: str,
        format: str,
        filename: str
    ) -> str:
        """
        导出报告

        Args:
            content: 报告内容
            format: 输出格式 (markdown, word, pdf)
            filename: 文件名

        Returns:
            文件路径
        """
        if format == "markdown":
            return await self._export_markdown(content, filename)
        elif format == "word":
            return await self._export_word(content, filename)
        elif format == "pdf":
            return await self._export_pdf(content, filename)
        else:
            raise ValueError(f"不支持的格式: {format}")

    async def _export_markdown(self, content: str, filename: str) -> str:
        """导出为Markdown格式"""
        file_path = os.path.join(self.export_dir, f"{filename}.md")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        return file_path

    async def _export_word(self, content: str, filename: str) -> str:
        """导出为Word格式"""
        import re

        file_path = os.path.join(self.export_dir, f"{filename}.docx")

        doc = Document()

        # 设置文档样式（宋体）
        from docx.oxml.ns import qn

        def set_style_font(style_obj, font_name='宋体'):
            """统一设置样式的中英文字体"""
            style_obj.font.name = font_name
            rpr = style_obj.element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = style_obj.element.makeelement(qn('w:rFonts'), {})
                rpr.insert(0, rfonts)
            rfonts.set(qn('w:eastAsia'), font_name)
            rfonts.set(qn('w:ascii'), font_name)
            rfonts.set(qn('w:hAnsi'), font_name)

        # Normal 样式
        set_style_font(doc.styles['Normal'])
        doc.styles['Normal'].font.size = Pt(12)

        # 标题样式全部设为宋体
        for heading_style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Title']:
            if heading_style_name in doc.styles:
                set_style_font(doc.styles[heading_style_name])

        def add_formatted_text(paragraph, text: str):
            """解析Markdown内联格式并添加到段落"""
            # 处理链接 [text](url) -> 显示实际URL
            text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\2', text)
            # 按 **粗体** 和 *斜体* 分段
            parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part)

        def add_heading_with_font(doc, text, level):
            """添加标题并强制设置宋体"""
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = '宋体'
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = run._element.makeelement(qn('w:rFonts'), {})
                    rpr.insert(0, rfonts)
                rfonts.set(qn('w:eastAsia'), '宋体')
                rfonts.set(qn('w:ascii'), '宋体')
                rfonts.set(qn('w:hAnsi'), '宋体')
            return heading

        lines = content.split("\n")

        for line in lines:
            # 标题
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line.lstrip("#").strip()
                title = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', title)

                if level == 1:
                    add_heading_with_font(doc, title, 0)
                elif level == 2:
                    add_heading_with_font(doc, title, 1)
                elif level == 3:
                    add_heading_with_font(doc, title, 2)
                else:
                    add_heading_with_font(doc, title, 3)

            # 无序列表（支持缩进）
            elif re.match(r'^[\s]*[-*]\s+', line):
                indent = len(line) - len(line.lstrip())
                text = re.sub(r'^[\s]*[-*]\s+', '', line)
                p = doc.add_paragraph(style='List Bullet')
                if indent >= 4:
                    p.paragraph_format.left_indent = Pt(36)
                add_formatted_text(p, text)

            # 有序列表 — 保留原始编号作为文本，不用 Word 自动编号
            elif re.match(r'^[\s]*\d+\.\s+', line):
                indent = len(line) - len(line.lstrip())
                m = re.match(r'^(\s*\d+\.\s+)(.*)', line)
                number = m.group(1)  # "1. " 或 "  17. "
                text = m.group(2)
                p = doc.add_paragraph()
                if indent >= 4:
                    p.paragraph_format.left_indent = Pt(36)
                # 编号部分加粗
                run_num = p.add_run(number)
                run_num.bold = True
                # 正文部分带格式
                add_formatted_text(p, text)

            # 空行
            elif line.strip() == "":
                continue

            # 普通文本
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, line)

        doc.save(file_path)
        return file_path

    async def _export_pdf(self, content: str, filename: str) -> str:
        """导出为PDF格式"""
        import re
        from xml.sax.saxutils import escape as xml_escape

        file_path = os.path.join(self.export_dir, f"{filename}.pdf")

        # 注册中文字体（宋体优先）
        font_name = 'Helvetica'
        try:
            import glob as _glob

            # 字体路径列表（按优先级排序：宋体 → 黑体 → 微软雅黑 → 其他）
            font_paths = [
                # 项目内置字体（最高优先级）
                os.path.join(os.path.dirname(__file__), "..", "fonts", "simsun.ttc"),
                # Windows 宋体
                "C:/Windows/Fonts/simsun.ttc",
                "C:/Windows/Fonts/simsun.ttc",
                # Windows 黑体/微软雅黑
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/msyh.ttc",
                "C:/Windows/Fonts/msyhbd.ttc",
                # Linux 宋体（手动安装路径）
                "/usr/share/fonts/truetype/simsun.ttc",
                "/usr/share/fonts/simsun.ttc",
                "/usr/local/share/fonts/simsun.ttc",
                # Linux 其他中文字体
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            ]

            # 动态搜索：扫描常见字体目录下所有 CJK/Noto/WQY 字体
            search_dirs = [
                "/usr/share/fonts/**/*CJK*",
                "/usr/share/fonts/**/*cjk*",
                "/usr/share/fonts/**/*wqy*",
                "/usr/share/fonts/**/*Noto*Sans*",
                "/usr/share/fonts/**/*noto*",
            ]
            for pattern in search_dirs:
                font_paths.extend(_glob.glob(pattern, recursive=True))

            for font_path in font_paths:
                if not os.path.exists(font_path):
                    continue
                try:
                    pdfmetrics.registerFont(TTFont('Chinese', font_path))
                    font_name = 'Chinese'
                    break
                except Exception:
                    # .ttc 文件可能需要 subfont index，尝试 index=0
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path, subfontIndex=0))
                        font_name = 'Chinese'
                        break
                    except Exception:
                        continue
        except Exception:
            pass

        # 创建PDF文档
        doc = SimpleDocTemplate(
            file_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # 定义样式
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=24,
            leading=30,
            alignment=1
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=18,
            leading=22,
            spaceAfter=12
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=14,
            spaceAfter=6
        )

        def sanitize_for_pdf(text: str) -> str:
            """清理Markdown文本，转为PDF安全的XML格式"""
            # 1. 处理Markdown链接 [text](url) -> 显示实际URL
            text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'\2', text)
            # 2. 处理粗体 **text** -> <b>text</b>
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            # 3. 处理斜体 *text* -> <i>text</i>（避免匹配已处理的<b>标签）
            text = re.sub(r'(?<!<b>)\*([^*]+?)\*(?!</b>)', r'<i>\1</i>', text)
            # 4. 转义XML特殊字符（但保留已插入的<b><i>标签）
            # 先保护标签
            placeholders = {}
            def protect_tag(m):
                key = f"__TAG{len(placeholders)}__"
                placeholders[key] = m.group(0)
                return key
            text = re.sub(r'</?[bi]>', protect_tag, text)
            # 转义
            text = xml_escape(text)
            # 恢复标签
            for key, val in placeholders.items():
                text = text.replace(key, val)
            return text

        # 解析内容并创建PDF元素
        story = []
        lines = content.split("\n")

        for line in lines:
            if line.startswith("# "):
                story.append(Paragraph(xml_escape(line[2:].strip()), title_style))
                story.append(Spacer(1, 12))

            elif line.startswith("## "):
                story.append(Paragraph(xml_escape(line[3:].strip()), heading_style))

            elif line.startswith("### "):
                story.append(Paragraph(xml_escape(line[4:].strip()), heading_style))

            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph("• " + sanitize_for_pdf(line[2:].strip()), body_style))

            elif line.strip():
                story.append(Paragraph(sanitize_for_pdf(line), body_style))

            else:
                story.append(Spacer(1, 12))

        # 生成PDF
        doc.build(story)

        return file_path

    def get_export_path(self, filename: str, format: str) -> str:
        """获取导出文件路径"""
        return os.path.join(self.export_dir, f"{filename}.{format}")
