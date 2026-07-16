"""
文档智能体 - 读取/分析/提取/总结文档（含 OCR）
"""
from __future__ import annotations

import logging
import os
import re
from typing import Optional, List, Dict, Any

from .base import Agent, Task

logger = logging.getLogger(__name__)


class DocumentAgent(Agent):
    """文档智能体"""

    # 图片文件扩展名
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}

    def __init__(self, llm: Optional[LLM] = None):
        super().__init__("document_agent", llm)
        self._ocr_model = None

    async def execute(self, task: Task) -> Dict[str, Any]:
        """执行文档处理任务"""
        file_path = task.input.get("file_path", "")
        file_type = task.input.get("file_type", "")
        action = task.input.get("action", "analyze")

        logger.info(f"[DocumentAgent] 收到任务 | 文件: {file_path} | 类型: {file_type} | 操作: {action}")

        if action == "analyze":
            return await self._analyze_document(file_path, file_type)
        elif action == "extract":
            return await self._extract_content(file_path, file_type)
        elif action == "summarize":
            return await self._summarize_document(file_path, file_type)
        else:
            logger.error(f"[DocumentAgent] 不支持的操作: {action}")
            raise ValueError(f"不支持的操作: {action}")

    def _is_image(self, file_path: str) -> bool:
        """判断文件是否为图片"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.IMAGE_EXTENSIONS

    def _get_ocr_model(self):
        """延迟加载 OCR 模型（RapidOCR + PP-OCRv6 ONNX）"""
        if self._ocr_model is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
                model_dir = os.path.join(os.path.dirname(__file__), "../models/PP-OCRv6")
                det_path = os.path.join(model_dir, "det_model/det.onnx")
                rec_path = os.path.join(model_dir, "rec.onnx")
                keys_path = os.path.join(model_dir, "ppocr_keys_v1.txt")
                logger.info(f"[DocumentAgent] 正在加载 RapidOCR + PP-OCRv6 模型")
                self._ocr_model = RapidOCR(
                    det_model_path=det_path,
                    rec_model_path=rec_path,
                    rec_keys_path=keys_path,
                )
                logger.info("[DocumentAgent] RapidOCR + PP-OCRv6 模型加载成功")
            except Exception as e:
                logger.error(f"[DocumentAgent] OCR 模型加载失败: {e}")
                raise RuntimeError(f"OCR 模型加载失败: {e}") from e
        return self._ocr_model

    async def _ocr_image(self, file_path: str) -> str:
        """使用 RapidOCR 识别图片中的文字"""
        logger.info(f"[DocumentAgent] 开始 OCR 识别 | 文件: {file_path}")

        if not os.path.exists(file_path):
            logger.error(f"[DocumentAgent] 图片文件不存在: {file_path}")
            raise FileNotFoundError(f"图片文件不存在: {file_path}")

        file_size = os.path.getsize(file_path)
        logger.info(f"[DocumentAgent] 图片大小: {file_size / 1024:.1f} KB")

        ocr = self._get_ocr_model()

        import time as _time
        t0 = _time.time()
        result, elapse = ocr(file_path)
        elapsed = _time.time() - t0

        # RapidOCR 返回 [(bbox, text, score), ...] 或 None
        texts = []
        if result:
            for item in result:
                try:
                    if isinstance(item, (list, tuple)) and len(item) >= 2:
                        text = item[1]
                        if text and str(text).strip():
                            texts.append(str(text).strip())
                except Exception as e:
                    logger.warning(f"[DocumentAgent] OCR 结果解析异常: {e}, item: {item}")
                    continue

        full_text = "\n".join(texts)
        logger.info(f"[DocumentAgent] OCR 识别完成 | 耗时: {elapsed:.2f}s | 识别到 {len(texts)} 条文本 | 总字符数: {len(full_text)}")
        return full_text

    async def _ocr_pdf(self, file_path: str) -> str:
        """将 PDF 逐页转图片后 OCR 识别（纯图片版 PDF）"""
        import time as _time
        import fitz  # PyMuPDF

        logger.info(f"[DocumentAgent] PDF 整页 OCR 模式 | 文件: {file_path}")
        t0 = _time.time()

        doc = fitz.open(file_path)
        total_pages = len(doc)
        logger.info(f"[DocumentAgent] PDF 共 {total_pages} 页，逐页转图片中...")

        all_texts = []
        for page_idx in range(total_pages):
            page = doc[page_idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_path = os.path.join(os.path.dirname(file_path), f"_pdf_ocr_page_{page_idx}.png")
            pix.save(img_path)

            try:
                page_text = await self._ocr_image(img_path)
                if page_text.strip():
                    all_texts.append(f"--- 第 {page_idx + 1} 页 ---\n{page_text}")
                    logger.info(f"[DocumentAgent] 第 {page_idx + 1}/{total_pages} 页 OCR 完成 | {len(page_text)} 字符")
                else:
                    logger.info(f"[DocumentAgent] 第 {page_idx + 1} 页无文字")
            finally:
                if os.path.exists(img_path):
                    os.remove(img_path)

        doc.close()
        full_text = "\n\n".join(all_texts)
        elapsed = _time.time() - t0
        logger.info(f"[DocumentAgent] PDF 整页 OCR 完成 | 耗时: {elapsed:.2f}s | 共 {len(all_texts)} 页有文字 | 总字符数: {len(full_text)}")
        return full_text

    async def _read_pdf_mixed(self, file_path: str) -> str:
        """混合 PDF：提取文字层 + 提取嵌入图片 OCR，合并结果"""
        import time as _time
        import fitz

        logger.info(f"[DocumentAgent] PDF 混合模式（文字+图片）| 文件: {file_path}")
        t0 = _time.time()

        doc = fitz.open(file_path)
        total_pages = len(doc)
        all_parts = []

        for page_idx in range(total_pages):
            page = doc[page_idx]
            page_parts = []

            # 1. 提取文字层
            text = page.get_text().strip()
            if text:
                page_parts.append(text)
                logger.info(f"[DocumentAgent] 第 {page_idx + 1} 页文字层: {len(text)} 字符")

            # 2. 提取嵌入图片并 OCR
            images = page.get_images(full=True)
            if images:
                logger.info(f"[DocumentAgent] 第 {page_idx + 1} 页发现 {len(images)} 张图片，正在 OCR...")
                for img_idx, img_info in enumerate(images):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        img_bytes = base_image["image"]
                        img_ext = base_image.get("ext", "png")
                        img_path = os.path.join(
                            os.path.dirname(file_path),
                            f"_pdf_img_p{page_idx}_{img_idx}.{img_ext}"
                        )
                        with open(img_path, "wb") as f:
                            f.write(img_bytes)

                        # 跳过过小的图片（图标、装饰等）
                        img_size = os.path.getsize(img_path)
                        if img_size < 5000:
                            os.remove(img_path)
                            continue

                        try:
                            ocr_text = await self._ocr_image(img_path)
                            if ocr_text.strip():
                                page_parts.append(f"[图片 {img_idx + 1} OCR]:\n{ocr_text}")
                                logger.info(f"[DocumentAgent] 第 {page_idx + 1} 页图片 {img_idx + 1} OCR: {len(ocr_text)} 字符")
                        finally:
                            if os.path.exists(img_path):
                                os.remove(img_path)
                    except Exception as e:
                        logger.warning(f"[DocumentAgent] 第 {page_idx + 1} 页图片 {img_idx + 1} 提取失败: {e}")

            # 3. 如果文字层为空且无图片 OCR 结果，整页转图片 OCR
            if not page_parts:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_path = os.path.join(os.path.dirname(file_path), f"_pdf_fallback_p{page_idx}.png")
                pix.save(img_path)
                try:
                    ocr_text = await self._ocr_image(img_path)
                    if ocr_text.strip():
                        page_parts.append(ocr_text)
                finally:
                    if os.path.exists(img_path):
                        os.remove(img_path)

            if page_parts:
                all_parts.append(f"--- 第 {page_idx + 1} 页 ---\n" + "\n\n".join(page_parts))

        doc.close()
        full_text = "\n\n".join(all_parts)
        elapsed = _time.time() - t0
        logger.info(f"[DocumentAgent] PDF 混合模式完成 | 耗时: {elapsed:.2f}s | 总字符数: {len(full_text)}")
        return full_text

    async def _read_word_mixed(self, file_path: str) -> str:
        """混合 Word：提取段落/表格 + 提取嵌入图片 OCR，合并结果"""
        import time as _time
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        logger.info(f"[DocumentAgent] Word 混合模式（文字+图片）| 文件: {file_path}")
        t0 = _time.time()

        doc = Document(file_path)
        text_parts = []

        # 1. 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading', '')
                    text_parts.append(f"{'#' * int(level)} {para.text}")
                else:
                    text_parts.append(para.text)

        # 2. 提取表格
        for i, table in enumerate(doc.tables):
            text_parts.append(f"\n--- 表格 {i + 1} ---")
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text_parts.append(row_text)

        text_content = "\n".join(text_parts)
        logger.info(f"[DocumentAgent] Word 文字提取: {len(text_content)} 字符")

        # 3. 提取嵌入图片并 OCR
        image_parts = []
        img_idx = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_idx += 1
                try:
                    img_data = rel.target_part.blob
                    img_ext = rel.target_part.content_type.split("/")[-1]
                    if img_ext == "jpeg":
                        img_ext = "jpg"
                    img_path = os.path.join(
                        os.path.dirname(file_path),
                        f"_word_img_{img_idx}.{img_ext}"
                    )
                    with open(img_path, "wb") as f:
                        f.write(img_data)

                    # 跳过过小的图片
                    if os.path.getsize(img_path) < 5000:
                        os.remove(img_path)
                        continue

                    try:
                        ocr_text = await self._ocr_image(img_path)
                        if ocr_text.strip():
                            image_parts.append(f"[图片 {img_idx} OCR]:\n{ocr_text}")
                            logger.info(f"[DocumentAgent] Word 图片 {img_idx} OCR: {len(ocr_text)} 字符")
                    finally:
                        if os.path.exists(img_path):
                            os.remove(img_path)
                except Exception as e:
                    logger.warning(f"[DocumentAgent] Word 图片 {img_idx} 提取失败: {e}")

        if image_parts:
            logger.info(f"[DocumentAgent] Word 共提取 {len(image_parts)} 张图片的 OCR 结果")

        # 4. 合并文字 + 图片 OCR
        parts = []
        if text_content.strip():
            parts.append(text_content)
        if image_parts:
            parts.append("\n\n--- 文档内图片识别 ---\n" + "\n\n".join(image_parts))

        full_text = "\n\n".join(parts)
        elapsed = _time.time() - t0
        logger.info(f"[DocumentAgent] Word 混合模式完成 | 耗时: {elapsed:.2f}s | 总字符数: {len(full_text)}")
        return full_text

    async def _read_document(self, file_path: str, file_type: str) -> str:
        """读取文档内容，自动判断是否需要 OCR"""
        if not os.path.exists(file_path):
            logger.error(f"[DocumentAgent] 文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_size = os.path.getsize(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"[DocumentAgent] 读取文档 | 文件: {os.path.basename(file_path)} | 扩展名: {file_ext} | 大小: {file_size / 1024:.1f} KB")

        # 图片文件 → 直接 OCR
        if self._is_image(file_path):
            logger.info(f"[DocumentAgent] 检测到图片文件，使用 OCR 识别 | 文件: {file_path}")
            content = await self._ocr_image(file_path)
            logger.info(f"[DocumentAgent] OCR 读取完成 | 内容长度: {len(content)} 字符")
            return content

        # PDF → 混合模式（文字层 + 嵌入图片 OCR）
        if file_ext == ".pdf":
            content = await self._read_pdf_mixed(file_path)
            logger.info(f"[DocumentAgent] PDF 混合读取完成 | 内容长度: {len(content)} 字符")
            return content

        # Word → 混合模式（段落 + 嵌入图片 OCR）
        if file_ext in (".docx", ".doc"):
            content = await self._read_word_mixed(file_path)
            logger.info(f"[DocumentAgent] Word 混合读取完成 | 内容长度: {len(content)} 字符")
            return content

        # 其他文件 → FileService 读取
        logger.info(f"[DocumentAgent] 使用 FileService 读取文档 | 类型: {file_type}")
        from services.file_service import FileService
        file_svc = FileService()
        content = await file_svc.read_file(file_path)
        logger.info(f"[DocumentAgent] 文档读取完成 | 内容长度: {len(content)} 字符")
        return content

    async def _analyze_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """分析文档"""
        logger.info(f"[DocumentAgent] 开始分析文档 | 文件: {file_path}")
        content = await self._read_document(file_path, file_type)

        prompt = f"""请分析以下文档内容，提取关键信息：

{content}

请提供：
1. 文档主题
2. 主要观点
3. 关键数据/事实
4. 文档结构"""

        analysis = await self.think(prompt)
        logger.info(f"[DocumentAgent] 文档分析完成 | 分析结果长度: {len(analysis)} 字符")

        return {
            "file_path": file_path,
            "file_type": file_type,
            "content_length": len(content),
            "analysis": analysis
        }

    async def _extract_content(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """提取文档内容"""
        logger.info(f"[DocumentAgent] 开始提取内容 | 文件: {file_path}")
        content = await self._read_document(file_path, file_type)
        logger.info(f"[DocumentAgent] 内容提取完成 | 长度: {len(content)} 字符")

        return {
            "file_path": file_path,
            "file_type": file_type,
            "content": content
        }

    async def _summarize_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """总结文档"""
        logger.info(f"[DocumentAgent] 开始总结文档 | 文件: {file_path}")
        content = await self._read_document(file_path, file_type)

        prompt = f"""请为以下文档生成一个简洁的摘要：

{content}

摘要应该包含主要观点和结论。"""

        summary = await self.think(prompt)
        logger.info(f"[DocumentAgent] 文档总结完成 | 摘要长度: {len(summary)} 字符")

        return {
            "file_path": file_path,
            "file_type": file_type,
            "summary": summary
        }
